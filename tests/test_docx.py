"""Guards on DOCX ingestion.

The risk with this format is not that extraction fails loudly. It is that a
DOCX has no coordinates, every layout rule in the package is built to measure
some, and a reader that invents plausible ones would produce confident output
derived from nothing. So most of these tests are about what must *not* happen:
no gutters found on a document with no pages, no blocks joined by a gap that
means nothing, no marker added to a line that already had one.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from restruct.document.stats import measure
from restruct.ingestion.docx import (
    ReflowableRenderer,
    is_heading_style,
    is_list_style,
    read_docx,
)
from tests.helpers import SYNTHETIC_DIRECTORY

FIXTURE = SYNTHETIC_DIRECTORY / "10.tight.docx"

needs_fixture = pytest.mark.skipif(
    not FIXTURE.exists(),
    reason="no DOCX fixture present",
)


def build(paragraphs: list[tuple[str, str]], tmp_path: Path) -> Path:
    """A minimal DOCX of (style, text) paragraphs."""
    import docx

    document = docx.Document()
    for style, text in paragraphs:
        document.add_paragraph(text, style=style)
    path = tmp_path / "built.docx"
    document.save(str(path))
    return path


# -- no geometry, and saying so ---------------------------------------------


@needs_fixture
def test_a_docx_declares_that_it_has_no_geometry() -> None:
    """Everything else here follows from this flag being honest."""
    assert read_docx(FIXTURE).has_geometry is False
    assert measure(read_docx(FIXTURE)).has_geometry is False


@needs_fixture
def test_no_column_gutters_are_found_in_a_document_with_no_pages() -> None:
    """A gutter is a corridor on a page. Ordinal boxes would happily yield
    one, and it would mean nothing."""
    assert measure(read_docx(FIXTURE)).column_gutters == ()


@needs_fixture
def test_nothing_continues_anything_across_a_meaningless_gap() -> None:
    """Consecutive ordinal boxes touch, so every gap measures zero and every
    block would join every other. A DOCX states its paragraph boundaries, so
    there is nothing to infer."""
    statistics = measure(read_docx(FIXTURE))
    assert not statistics.is_paragraph_gap(0.0)
    assert not statistics.is_paragraph_gap(1.0)
    assert statistics.cell_gap_threshold == 0.0


@needs_fixture
def test_reading_order_is_preserved(tmp_path: Path) -> None:
    document = read_docx(FIXTURE)
    tops = [line.bbox[1] for line in document.lines]
    assert tops == sorted(tops)
    assert len(set(tops)) == len(tops), "two lines share an ordinal position"


# -- what the document states outright --------------------------------------


@pytest.mark.parametrize(
    ("style", "heading"),
    [("Heading 1", True), ("Heading 4", True), ("Title", True), ("Normal", False)],
)
def test_heading_styles_are_recognised(style: str, heading: bool) -> None:
    assert is_heading_style(style) is heading


@pytest.mark.parametrize(
    ("style", "listed"),
    [("List Bullet", True), ("List Number", True), ("Normal", False)],
)
def test_list_styles_are_recognised(style: str, listed: bool) -> None:
    assert is_list_style(style) is listed


def test_a_list_style_gets_the_marker_the_document_renders(tmp_path: Path) -> None:
    """Word draws the marker from the numbering part, not the run text. Putting
    it back is what lets nine downstream bullet rules stay unchanged."""
    path = build([("List Bullet", "Built pipelines.")], tmp_path)
    assert read_docx(path).lines[0].text == "• Built pipelines."


def test_a_marker_the_document_already_wrote_is_not_doubled(tmp_path: Path) -> None:
    """Testing the text after the fact cannot tell a marker this reader added
    from one that was always there, and doubles the second kind."""
    path = build([("List Bullet", "• Already marked")], tmp_path)
    assert read_docx(path).lines[0].text == "• Already marked"


def test_a_style_size_reaches_the_span(tmp_path: Path) -> None:
    """Word records a heading's size on the style, not on the run. Reading the
    run alone reports 0.0 for everything and leaves nothing to compare."""
    path = build([("Heading 1", "WORK EXPERIENCE"), ("Normal", "Body text")], tmp_path)
    lines = read_docx(path).lines
    assert lines[0].spans[0].size > lines[1].spans[0].size > 0
    assert lines[0].spans[0].bold


def test_a_table_states_its_own_cells(tmp_path: Path) -> None:
    """No gap measurement is involved: the source says which cell is which."""
    import docx

    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Training"
    table.cell(0, 1).text = "Year"
    table.cell(1, 0).text = "Safety"
    table.cell(1, 1).text = "2024"
    path = tmp_path / "table.docx"
    document.save(str(path))

    cells = {line.text: line.table_cell for line in read_docx(path).lines}
    assert cells["Training"][1:] == (0, 0)
    assert cells["Year"][1:] == (0, 1)
    assert cells["Safety"][1:] == (1, 0)


def test_stated_table_rows_group_without_measuring_anything(tmp_path: Path) -> None:
    from restruct.ingestion.native import extracted_lines
    from restruct.layout.rows import _visual_rows

    import docx

    document = docx.Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Training"
    table.cell(0, 1).text = "Year"
    table.cell(1, 0).text = "Safety"
    table.cell(1, 1).text = "2024"
    path = tmp_path / "rows.docx"
    document.save(str(path))

    physical = read_docx(path)
    statistics = measure(physical)
    lines = extracted_lines(physical, statistics)
    rows = _visual_rows(lines, range(len(lines)), statistics)
    assert [[line.text for _, line in row] for row in rows] == [
        ["Training", "Year"],
        ["Safety", "2024"],
    ]


def test_a_paragraph_between_two_tables_keeps_its_place(tmp_path: Path) -> None:
    """python-docx lists paragraphs and tables separately, which loses the
    interleaving. Walking the body element is what keeps document order."""
    import docx

    document = docx.Document()
    document.add_paragraph("before")
    document.add_table(rows=1, cols=1).cell(0, 0).text = "inside"
    document.add_paragraph("after")
    path = tmp_path / "order.docx"
    document.save(str(path))
    assert [line.text for line in read_docx(path).lines] == [
        "before",
        "inside",
        "after",
    ]


# -- the renderer stand-in ---------------------------------------------------


@needs_fixture
def test_the_renderer_answers_what_the_parsers_ask() -> None:
    """Every parser takes a renderer. Rather than a DOCX branch in each of
    them, the stand-in answers "no glyphs here" and lets span resolution fall
    back to proportional estimation, which it already does for OCR."""
    document = read_docx(FIXTURE)
    renderer = ReflowableRenderer(document)
    page = renderer[0]
    assert page.search_for("anything") == []
    assert page.rect.height > 0
    assert isinstance(page.get_text("words"), list)


# -- the pass-1 dump says what the DOCX said ---------------------------------


def test_the_pass_one_dump_records_what_python_docx_read(tmp_path: Path) -> None:
    """A DOCX dumped in PyMuPDF's shape used to be an empty page list: nothing
    to reason from. The dump holds the facts this reader worked off instead --
    style names, indent, cells, resolved fonts -- and not the ordinal boxes,
    which are the one thing here that was invented rather than read."""
    path = build([("Heading 1", "EXPERIENCE"), ("List Bullet", "Led it")], tmp_path)
    page = read_docx(path).raw_pages[0]

    assert page["reader"] == "python-docx"
    assert page["defaultFontSize"] > 0
    heading, bullet = page["paragraphs"]
    assert heading["style"] == "Heading 1" and heading["isHeadingStyle"]
    assert bullet["isListStyle"] and bullet["listMarkerAdded"]
    assert bullet["text"] == "• Led it"
    assert heading["runs"][0]["bold"] and heading["runs"][0]["size"] > 0
    assert "bbox" not in heading


def test_the_dump_separates_what_the_run_said_from_what_it_inherited(
    tmp_path: Path,
) -> None:
    """The resolved value is the same either way, and only one of them is
    evidence about this line rather than about every line sharing its style."""
    import docx

    document = docx.Document()
    document.add_paragraph("HEADING", style="Heading 1")
    document.add_paragraph().add_run("Emphasised").bold = True
    path = tmp_path / "stated.docx"
    document.save(str(path))

    inherited, direct = read_docx(path).raw_pages[0]["paragraphs"]
    assert inherited["runs"][0]["bold"] and "statedOnRun" not in inherited["runs"][0]
    assert direct["runs"][0]["statedOnRun"] == ["bold"]


def test_the_dump_locates_a_table_cell_by_what_the_table_stated(
    tmp_path: Path,
) -> None:
    import docx

    document = docx.Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Training"
    table.cell(0, 1).text = "2024"
    path = tmp_path / "table.docx"
    document.save(str(path))

    paragraphs = read_docx(path).raw_pages[0]["paragraphs"]
    assert [entry["tableCell"][1:] for entry in paragraphs] == [[0, 0], [0, 1]]


def test_the_dump_is_named_after_the_reader_that_produced_it() -> None:
    """The two readers answer different questions, so a DOCX dump named
    ``raw-pymupdf.json`` would misdescribe its own contents."""
    from restruct.stages import raw_extraction_reader

    assert raw_extraction_reader(Path("resume.docx")) == "docx"
    assert raw_extraction_reader(Path("RESUME.DOCX")) == "docx"
    assert raw_extraction_reader(Path("resume.pdf")) == "pymupdf"


def test_an_unreadable_file_is_reported_as_an_invalid_document(tmp_path: Path) -> None:
    from restruct.errors import InvalidDocument

    broken = tmp_path / "broken.docx"
    broken.write_bytes(b"not a zip at all")
    with pytest.raises(InvalidDocument):
        read_docx(broken)
