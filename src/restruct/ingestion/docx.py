"""DOCX ingestion, converging on the same physical types as PDF.

A DOCX is a zip of XML describing what the author wrote, not what a renderer
drew. There are no coordinates anywhere in it, and that is the whole design
problem: every layout rule in this package exists to *infer* from geometry
something a PDF declined to say.

MuPDF will open a DOCX and hand back convincing bounding boxes. It gets them by
re-laying the document out with substituted fonts into an invented page box, and
in doing so loses the paragraph styles, the list markers and the table
structure. Those boxes would then feed every document-relative statistic we
have -- gutters, baselines, cell gaps, paragraph thresholds -- which would
measure a layout nobody ever laid out and produce plausible output from
fiction. Reading the XML is slower to write and says only true things.

Losing geometry costs less than it sounds, because the DOCX states outright
almost everything the geometric heuristics were reconstructing:

    heading           style name, not "larger and bolder than the body"
    bullet            style name and numbering, not a regex over marker glyphs
    line boundaries   one paragraph is one line, not a vertical-gap threshold
    table cells       real rows and cells, not a wide-gap heuristic

What is left in ``bbox`` is ordinal: reading-order position and indent depth,
so ordering and nesting survive intact. ``Document.has_geometry`` is False, and
that is what stops any threshold reading points into numbers that carry none.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterator

from dataclasses import replace

from restruct.document.physical import Document, Page, Span, TextLine, Token
from restruct.patterns.bullets import BULLET_RE
from restruct.errors import InvalidDocument

# One paragraph occupies one unit of the ordinal y axis. The value is arbitrary
# and deliberately so: nothing may read a measurement out of it, and a round
# number makes that obvious in a debug dump.
_LINE_HEIGHT = 1.0

# How far one level of indentation moves a line along the ordinal x axis.
_INDENT_STEP = 1.0

# A page has to have some width for the boxes to be non-degenerate. This is a
# count of indent steps, not a page size.
_ORDINAL_WIDTH = 100.0

# PyMuPDF span flags, reproduced so a DOCX run and a PDF span answer ``bold``
# and ``italic`` through exactly the same code path.
_FLAG_ITALIC = 1 << 1
_FLAG_BOLD = 1 << 4

# Word draws a marker for a list paragraph, but stores it in the numbering
# part rather than in the run text. Putting it back in the text is
# reconstruction, not invention: the rendered document has it, a PDF of the
# same file would have it in the text, and OCR would read it. Normalising here
# is what lets nine downstream bullet rules stay unchanged.
_LIST_MARKER = "\u2022 "

# Word's own names for the styles that mean something to us. Matched
# case-insensitively by prefix, so "Heading 1" and "heading 2" both resolve and
# a theme that renames "List Bullet" to "List Bullet 2" still reads as a list.
_HEADING_STYLE_PREFIX = "heading"
_TITLE_STYLE_NAMES = frozenset({"title", "subtitle"})
_LIST_STYLE_PREFIXES = ("list bullet", "list number", "list paragraph")


def _resolved_font_property(paragraph: Any, run: Any, name: str) -> Any:
    """Read a run property, following the style chain when the run inherits.

    python-docx reports None for "inherit", which is the common case: a resume
    written properly sets bold once on the Heading style rather than on every
    run. Reading the run alone would report almost nothing as bold.
    """
    direct = getattr(run.font, name, None)
    if direct is not None:
        return direct
    # Both chains, in this order. The run's character style is almost always
    # "Default Paragraph Font", which states nothing and is nonetheless not
    # None -- treating it as the only fallback silently skips the paragraph
    # style, which is where Word actually records that a heading is bold.
    for start in (getattr(run, "style", None), paragraph.style):
        style = start
        seen = 0
        while style is not None and seen < 10:  # styles can, in theory, cycle
            value = getattr(style.font, name, None)
            if value is not None:
                return value
            style = getattr(style, "base_style", None)
            seen += 1
    return None


def _run_boxes(
    runs: list[Any],
    left: float,
    top: float,
) -> list[tuple[float, float, float, float]]:
    """Lay runs along the ordinal x axis in proportion to their length.

    Every run of a paragraph sharing one box would leave word reconstruction
    with identical coordinates for every word, and would make a hyperlink
    indistinguishable from the sentence around it. Proportional slices are
    still ordinal -- they measure characters, not points -- but they order and
    separate, which is what the later passes actually ask of a box.
    """
    total = sum(len(run.text) for run in runs) or 1
    span_width = _ORDINAL_WIDTH - left
    boxes: list[tuple[float, float, float, float]] = []
    cursor = left
    for run in runs:
        width = span_width * (len(run.text) / total)
        boxes.append((cursor, top, cursor + width, top + _LINE_HEIGHT))
        cursor += width
    return boxes


_WORD_NAMESPACE = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _document_default_size(document: Any) -> float:
    """The body size from docDefaults, where Word records it.

    Styles that state no size inherit from here, and "Normal" usually does.
    Without it every body paragraph reports 0.0, the body size measures as
    zero, and "larger than the body" can never be true -- which would leave
    heading detection with nothing to compare against.
    """
    try:
        element = document.styles.element
    except AttributeError:
        return 0.0
    for path in (
        f"{_WORD_NAMESPACE}docDefaults/{_WORD_NAMESPACE}rPrDefault/"
        f"{_WORD_NAMESPACE}rPr/{_WORD_NAMESPACE}sz",
    ):
        found = element.find(path)
        if found is not None:
            value = found.get(f"{_WORD_NAMESPACE}val")
            if value and value.isdigit():
                # Word stores the size in half-points.
                return float(value) / 2.0
    return 0.0


def _span_from_run(
    paragraph: Any,
    run: Any,
    box: tuple[float, ...],
    default_size: float = 0.0,
) -> Span:
    bold = bool(_resolved_font_property(paragraph, run, "bold"))
    italic = bool(_resolved_font_property(paragraph, run, "italic"))
    size = _resolved_font_property(paragraph, run, "size")
    name = _resolved_font_property(paragraph, run, "name")
    flags = (_FLAG_BOLD if bold else 0) | (_FLAG_ITALIC if italic else 0)
    text = run.text
    return Span(
        text=text,
        bbox=tuple(float(value) for value in box),
        font=str(name or ""),
        # Points, from the document's own styles rather than from a renderer's
        # substitution. Zero when the style chain never states one.
        size=float(size.pt) if size is not None else default_size,
        flags=flags,
        # OCR-style granularity: the run is the smallest unit the source
        # reports, exactly as a Tesseract word is. Word reconstruction already
        # knows what to do with that, so nothing downstream needs a DOCX case.
        tokens=(Token(text=text, bbox=tuple(float(v) for v in box)),),
        granularity="word",
    )


def _style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    return str(getattr(style, "name", "") or "")


def is_heading_style(style: str) -> bool:
    """Whether a style name is Word's way of saying "this is a heading"."""
    folded = style.strip().casefold()
    return folded.startswith(_HEADING_STYLE_PREFIX) or folded in _TITLE_STYLE_NAMES


def is_list_style(style: str) -> bool:
    """Whether a style name is Word's way of saying "this is a list item"."""
    folded = style.strip().casefold()
    return folded.startswith(_LIST_STYLE_PREFIXES)


def _indent_level(paragraph: Any) -> float:
    """Indent depth in steps, from the paragraph's own left indent.

    Ordinal, not a measurement: it records that one paragraph sits inside
    another, which is all the nesting a reader can act on.
    """
    indent = getattr(paragraph.paragraph_format, "left_indent", None)
    if indent is None:
        return 0.0
    # Half an inch is Word's default list indent, so this counts list levels.
    return max(0.0, round(float(indent.inches) / 0.5, 2))


def _block_items(container: Any) -> Iterator[tuple[str, Any]]:
    """Paragraphs and tables in the order the document lays them out.

    python-docx exposes ``paragraphs`` and ``tables`` as separate lists, which
    loses the interleaving -- a table between two paragraphs would be read as
    both paragraphs and then the table. Walking the body element keeps the
    document's own order, which is the only reading order a DOCX has.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = container.element.body if hasattr(container, "element") else container._tc
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield "paragraph", Paragraph(child, container)
        elif child.tag.endswith("}tbl"):
            yield "table", Table(child, container)


def _lines_from_container(
    container: Any,
    cursor: list[float],
    links: list[Any],
    default_size: float = 0.0,
    table_path: tuple[int, int, int] | None = None,
    indent: float = 0.0,
) -> Iterator[TextLine]:
    """Every paragraph in a body or a table cell, in document order.

    ``links`` is filled as a side effect because ``TextLine`` is frozen and a
    link belongs to the page, not to the line -- the same place a PDF puts it.
    """
    for kind, item in _block_items(container):
        if kind == "paragraph":
            line = _line_from_paragraph(
                item, cursor, links, default_size, table_path, indent
            )
            if line is not None:
                yield line
            continue
        for row_index, row in enumerate(item.rows):
            for column_index, cell in enumerate(row.cells):
                yield from _lines_from_container(
                    cell,
                    cursor,
                    links,
                    default_size,
                    (int(cursor[1]), row_index, column_index),
                    indent + _INDENT_STEP,
                )
        cursor[1] += 1


def _line_from_paragraph(
    paragraph: Any,
    cursor: list[float],
    links: list[Any],
    default_size: float,
    table_path: tuple[int, int, int] | None,
    indent: float,
) -> TextLine | None:
    text = paragraph.text
    if not text.strip():
        return None

    style = _style_name(paragraph)
    # Only when the source did not write one itself. Testing the text after
    # the fact cannot tell a marker this reader added from one the document
    # always had, and doubles the second kind.
    marker_added = is_list_style(style) and not BULLET_RE.match(text)
    if marker_added:
        text = f"{_LIST_MARKER}{text}"

    left = (_indent_level(paragraph) + indent) * _INDENT_STEP
    top = cursor[0]
    cursor[0] += _LINE_HEIGHT
    box = (left, top, _ORDINAL_WIDTH, top + _LINE_HEIGHT)

    runs = [run for run in paragraph.runs if run.text]
    boxes = _run_boxes(runs, left, top)
    spans = tuple(
        _span_from_run(paragraph, run, run_box, default_size)
        for run, run_box in zip(runs, boxes, strict=True)
    )
    if spans and marker_added:
        first = spans[0]
        spans = (replace(first, text=f"{_LIST_MARKER}{first.text}"),) + spans[1:]
    if not spans:
        # A paragraph whose text lives outside any run -- a field result, a
        # hyperlink -- still has text worth keeping.
        spans = (
            Span(
                text=text,
                bbox=box,
                font="",
                size=0.0,
                flags=0,
                tokens=(Token(text=text, bbox=box),),
                granularity="word",
            ),
        )
    links.extend(_hyperlinks(paragraph, box))
    return TextLine(
        page=1,
        bbox=box,
        spans=spans,
        style=style,
        table_cell=table_path,
    )


def _hyperlinks(paragraph: Any, box: tuple[float, float, float, float]) -> list[Any]:
    """The links a paragraph carries, with their targets resolved.

    A DOCX states the target and the visible text exactly, so nothing has to be
    matched by position the way a PDF annotation does. The box is the
    paragraph's, which is all the location an unrendered link has.
    """
    from restruct.document.physical import LinkAnnotation

    found: list[Any] = []
    part = paragraph.part
    for element in paragraph._p.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
    ):
        relationship_id = element.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
        )
        if not relationship_id:
            continue
        try:
            target = part.rels[relationship_id].target_ref
        except KeyError:
            continue
        if target:
            found.append(LinkAnnotation(page=1, bbox=box, uri=str(target)))
    return found


def read_docx(path: Path) -> Document:
    """Read one DOCX into the shared physical representation.

    A DOCX is reflowable and has no pages of its own -- what Word shows depends
    on the printer it is aimed at -- so everything lands on one logical page.
    """
    try:
        import docx
    except ImportError as error:  # pragma: no cover - dependency is declared
        raise InvalidDocument(path, f"DOCX support is unavailable: {error}") from error

    try:
        document = docx.Document(str(path))
    except Exception as error:
        raise InvalidDocument(path, str(error)) from error

    cursor = [0.0, 0.0]
    collected_links: list[Any] = []
    lines = tuple(
        _lines_from_container(
            document, cursor, collected_links, _document_default_size(document)
        )
    )
    links = tuple(collected_links)
    page = Page(
        number=1,
        width=_ORDINAL_WIDTH,
        # Tall enough to hold every line, so nothing reads as off the page.
        height=max(cursor[0], _LINE_HEIGHT),
        lines=lines,
        links=links,
    )
    return Document(pages=(page,), has_geometry=False)


class _ReflowablePage:
    """Stands in for a rendered page that a reflowable source does not have.

    The section parsers ask the renderer two things: where a span sits on the
    page, and how tall the page is. Neither has an answer here, and the honest
    ones are already handled downstream -- ``resolve_span_box`` falls back to
    proportional estimation when a search finds nothing, so returning no hits
    is not a failure but a statement.
    """

    def __init__(self, height: float, page: Any = None) -> None:
        import pymupdf

        self.rect = pymupdf.Rect(0.0, 0.0, _ORDINAL_WIDTH, height)
        self._page = page

    def search_for(self, *arguments: Any, **keywords: Any) -> list:
        """No rendered glyphs to search. Says so rather than guessing."""
        return []

    def get_links(self) -> list[dict[str, Any]]:
        """The hyperlinks the XML stated, shaped the way PyMuPDF reports them.

        A DOCX names the target outright, so nothing is being recovered from a
        rendering here -- only reshaped, so the annotation matcher needs no
        DOCX case of its own.
        """
        if self._page is None:
            return []
        import pymupdf

        return [
            {"uri": link.uri, "from": pymupdf.Rect(link.bbox)}
            for link in self._page.links
        ]

    def get_text(self, kind: str, **keywords: Any) -> list:
        """Words with their ordinal boxes, in PyMuPDF's tuple shape.

        Only "words" is answered, because only the link matcher asks, and it
        matches word boxes against link boxes -- proportional run boxes are
        exactly enough for that.
        """
        if kind != "words" or self._page is None:
            return []
        found = []
        for line_number, line in enumerate(self._page.lines):
            for word_number, word in enumerate(line.words):
                found.append(
                    (*word.bbox, word.text, 0, line_number, word_number)
                )
        return found


class ReflowableRenderer:
    """A renderer-shaped object for a document that was never rendered.

    Keeps every parser signature unchanged rather than threading an optional
    renderer through nine call sites, which would put a DOCX branch in each of
    them.
    """

    def __init__(self, document: Document) -> None:
        self._pages = [_ReflowablePage(page.height, page) for page in document.pages]

    @property
    def page_count(self) -> int:
        return len(self._pages)

    def __getitem__(self, index: int) -> _ReflowablePage:
        return self._pages[index]

    def __len__(self) -> int:
        return len(self._pages)
